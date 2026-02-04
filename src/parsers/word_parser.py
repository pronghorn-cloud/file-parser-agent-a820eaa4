"""Word document parser using python-docx."""

from pathlib import Path
from typing import Any, Dict, List
import logging

try:
    from docx import Document
    from docx.table import Table
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .base import BaseDocumentParser, ParsedDocument

logger = logging.getLogger(__name__)


class WordParser(BaseDocumentParser):
    """Parser for Word documents.
    
    Extracts paragraphs, tables, styles, and metadata from
    Word documents using the python-docx library.
    """
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def __init__(self, file_path: Path):
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word parsing. Install with: pip install python-docx")
        super().__init__(file_path)
        self._document = None
        
        # Note: .doc files require additional handling
        if self.file_path.suffix.lower() == '.doc':
            logger.warning(".doc files may have limited support. Consider converting to .docx")
    
    @property
    def document(self) -> 'Document':
        """Lazy-load Word document."""
        if self._document is None:
            self._document = Document(str(self.file_path))
        return self._document
    
    def parse(self) -> ParsedDocument:
        """Parse Word document and return structured content."""
        logger.info(f"Parsing Word document: {self.file_path}")
        
        doc = ParsedDocument(
            filename=self.file_path.name,
            file_type='word'
        )
        
        try:
            doc.metadata = self.extract_metadata()
            doc.content = {
                'paragraphs': self._extract_paragraphs(),
                'sections': self._extract_sections()
            }
            doc.tables = self.extract_tables()
            doc.images = self.extract_images()
        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            doc.errors.append(str(e))
        
        return doc
    
    def extract_text(self) -> str:
        """Extract all text from Word document."""
        paragraphs = [p.text for p in self.document.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    
    def extract_metadata(self) -> Dict[str, Any]:
        """Extract Word document metadata."""
        metadata = {}
        core_props = self.document.core_properties
        
        # Extract available core properties
        prop_names = [
            'author', 'category', 'comments', 'content_status',
            'created', 'identifier', 'keywords', 'language',
            'last_modified_by', 'last_printed', 'modified',
            'revision', 'subject', 'title', 'version'
        ]
        
        for prop in prop_names:
            value = getattr(core_props, prop, None)
            if value:
                # Convert datetime objects to ISO format
                if hasattr(value, 'isoformat'):
                    value = value.isoformat()
                metadata[prop] = str(value)
        
        # Add document statistics
        metadata['paragraph_count'] = len(self.document.paragraphs)
        metadata['table_count'] = len(self.document.tables)
        metadata['section_count'] = len(self.document.sections)
        
        return metadata
    
    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """Extract paragraphs with styling information."""
        paragraphs = []
        for i, para in enumerate(self.document.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                para_data = {
                    'index': i,
                    'text': para.text,
                    'style': para.style.name if para.style else None,
                }
                
                # Check if it's a heading
                if para.style and para.style.name.startswith('Heading'):
                    para_data['is_heading'] = True
                    # Extract heading level
                    try:
                        level = int(para.style.name.replace('Heading ', ''))
                        para_data['heading_level'] = level
                    except ValueError:
                        para_data['heading_level'] = 1
                
                paragraphs.append(para_data)
        
        return paragraphs
    
    def _extract_sections(self) -> List[Dict[str, Any]]:
        """Extract section information."""
        sections = []
        for i, section in enumerate(self.document.sections):
            section_data = {
                'index': i,
                'page_width': section.page_width.inches if section.page_width else None,
                'page_height': section.page_height.inches if section.page_height else None,
                'orientation': 'landscape' if section.page_width and section.page_height 
                               and section.page_width > section.page_height else 'portrait',
                'left_margin': section.left_margin.inches if section.left_margin else None,
                'right_margin': section.right_margin.inches if section.right_margin else None,
                'top_margin': section.top_margin.inches if section.top_margin else None,
                'bottom_margin': section.bottom_margin.inches if section.bottom_margin else None,
            }
            sections.append(section_data)
        
        return sections
    
    def extract_tables(self) -> List[Dict[str, Any]]:
        """Extract tables from Word document."""
        tables = []
        for i, table in enumerate(self.document.tables):
            table_data = {
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'data': []
            }
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['data'].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def extract_images(self) -> List[Dict[str, Any]]:
        """Extract image references from Word document."""
        images = []
        
        # Access inline shapes (images)
        for i, shape in enumerate(self.document.inline_shapes):
            image_data = {
                'index': i,
                'width': shape.width.inches if shape.width else None,
                'height': shape.height.inches if shape.height else None,
                'type': str(shape.type)
            }
            images.append(image_data)
        
        return images
